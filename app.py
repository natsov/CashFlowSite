import base64
import secrets
import smtplib
from email.message import EmailMessage
from flask import Flask, flash, render_template, url_for, request, redirect, session, send_from_directory,make_response,jsonify
from datetime import datetime
from werkzeug.utils import secure_filename
import os
from PIL import Image
import io
import pandas as pd
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt
from db import CashDB
import bcrypt
matplotlib.use('Agg')
from flask_mail import Mail, Message
import random
import string
import smtplib
import jsonschema
from email.mime.text import MIMEText

CashDB = CashDB('cashflow.db')

app = Flask(__name__, static_folder='static')
app.config['UPLOAD_FOLDER'] = 'static/images_wishes'
app.secret_key = '1234'

def generate_password(length=8):
    """Генерирует случайный пароль заданной длины"""
    characters = string.ascii_letters + string.digits + string.punctuation
    password = ''.join(random.choice(characters) for _ in range(length))
    return password


@app.route('/forgot_password', methods=['GET', 'POST'])
def send_password_reset_email():
    """Отправляет письмо с новым паролем пользователю"""
    if request.method == 'POST':
        email = request.form['email']  # Получаем адрес электронной почты из формы
        new_password = generate_password()
        print(email, new_password)
        sender_email = 'app.cashflow.g@gmail.com'  # Замените на вашу почту
        sender_password = 'dcmlomsgqhojtztj'  # Замените на пароль от вашей почты
        if CashDB.check_email(email):
            subject = 'Сброс пароля'
            message = f'Ваш новый пароль: {new_password}'
            msg = MIMEText(message)
            msg['Subject'] = subject
            msg['From'] = sender_email
            msg['To'] = email
            try:
                with smtplib.SMTP('smtp.gmail.com', 587) as smtp:
                    smtp.starttls()
                    smtp.login(sender_email, sender_password)
                    smtp.send_message(msg)
                hashed_password = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt())
                password = hashed_password
                result = CashDB.update_password(email, password)
                session['email'] = email
                print(session['email'])
                return redirect("/reset_password")
            except smtplib.SMTPException as e:
                print('Ошибка при отправке письма:', str(e))
        else:
            error_message = "Ошибка при авторизации. Попробуйте ещё раз"
            return render_template("forgot_password.html", error_message=error_message)
    return render_template("forgot_password.html")

@app.route('/reset_password', methods=['GET', 'POST'])
def sign_in_new_password():
    if request.method == 'POST':
        email = session['email']
        print(email)
        password = request.form['password']
        stored_password = CashDB.get_password(email)
        if stored_password:
            hashed_password = stored_password[0][0].tobytes()
            if bcrypt.checkpw(password.encode('utf-8'), hashed_password):
                session['email'] = email
                return redirect(url_for('my_page'))
            else:
                error_message = "Ошибка при авторизации. Попробуйте ещё раз"
                return render_template("reset_password.html", error_message=error_message)
    return render_template("reset_password.html")

@app.route('/')
@app.route('/home')
def index():
    return render_template("index.html")


@app.route('/about')
def about():
    return render_template("about.html")


@app.route('/my-page', methods=['GET', 'POST'])
def my_page():
    if 'email' in session:
        email = session['email']
        result = CashDB.get_user_name(email)
        return render_template("my-page.html", name_surname=result)
    else:
        return redirect(url_for('sign_in'))


@app.route('/sign-in', methods=['POST', 'GET'])
def sign_in():
    error_message = None
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        stored_password = CashDB.get_password(email)
        admin_password = CashDB.get_password_admin(email)
        if stored_password:
            hashed_password = stored_password[0][0].tobytes()  # Преобразовать memoryview в bytes
            if bcrypt.checkpw(password.encode('utf-8'), hashed_password):
                session['email'] = email
                return redirect(url_for('my_page'))
            else:
                error_message = "Ошибка при авторизации. Попробуйте ещё раз"
        elif admin_password:
            hashed_password = admin_password[0][0].tobytes()
            if bcrypt.checkpw(password.encode('utf-8'), hashed_password):
                session['email'] = email
                return redirect('/admin-page')
            else:
                error_message = "Ошибка при авторизации. Попробуйте ещё раз"
    return render_template("sign-in.html", error_message=error_message)



@app.route('/sign-up', methods=['POST', 'GET'])
def sign_up():
    error_message = None
    if request.method == 'POST':
        name = request.form['name']
        second_name = request.form['second-name']
        email = request.form['email']
        password = request.form['password']
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        password = hashed_password
        result = CashDB.add_user(name, second_name, email, password)
        userid = CashDB.get_user_id(email)
        if result:
            session['email'] = email
            return redirect('/my-page')
        else:
            error_message = "Ошибка при регистрации. Попробуйте ввести другие данные"
    return render_template("sign-up.html", error_message=error_message)


@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        result = CashDB.check_email(email)
        if result:
            msg = Message('Сброс пароля', sender='cash_flow@gmail.com', recipients=[email])
            msg.body = 'Для сброса пароля перейдите по ссылке: {}'.format(url_for('forgot_password', _external=True))
            mail.send(msg)
            return redirect(url_for('password_reset_sent'))
    return render_template('forgot_password.html')

def get_users():
    with open("users.json", "r") as json_file:
        users = json.load(json_file)
    return users

def save_users(users):
    with open("users.json", "w") as json_file:
        json.dump(users, json_file)

@app.route('/debts', methods=['POST', 'GET'])
def debts():
    error_message = None
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    debts_user = CashDB.show_debts(userid[0][0])
    current_datetime = datetime.now()
    if request.method == 'POST':
        sort_by = request.form.get("sort_values")
        sort_order = request.form.get("sort_order")
        if sort_by == "sum":
            sorted_data = CashDB.order_by("debtamount", userid[0][0], sort_order)
            return render_template("debts.html", debts=sorted_data, current_datetime=current_datetime)
        elif sort_by == "alphabet":
            sorted_data = CashDB.order_by("creditorname", userid[0][0], sort_order)
            return render_template("debts.html", debts=sorted_data, current_datetime=current_datetime)
        elif sort_by == "date":
            sorted_data = CashDB.order_by("duedate", userid[0][0], sort_order)
            return render_template("debts.html", debts=sorted_data, current_datetime=current_datetime)
        else:
            sorted_data = None
        if not sort_by:
            if 'show_form' in request.form:
                show_form = True
                return render_template("debts.html", show_form=show_form)
            amount = request.form['amount']
            name = request.form['name']
            date = request.form['date']

            if not amount:
                show_form = True
                error_message = "Не введена сумма. Попробуйте снова"
                return render_template("debts.html", show_form=show_form, debts=debts_user, error_message=error_message,  current_datetime=current_datetime)
            elif float(amount) <= 0:
                show_form = True
                error_message = "Сумма не может быть отрицательной. Попробуйте снова"
                return render_template("debts.html", show_form=show_form, debts=debts_user, error_message=error_message,  current_datetime=current_datetime)

            elif not name:
                show_form = True
                error_message = "Не введено имя или название организации! Попробуйте снова"
                return render_template("debts.html", show_form=show_form, debts=debts_user, error_message=error_message,  current_datetime=current_datetime)

            elif not date:
                show_form = True
                error_message = "Не введена дата выполнения цели! Попробуйте снова"
                return render_template("debts.html", show_form=show_form, debts=debts_user, error_message=error_message,  current_datetime=current_datetime)
            else:
                result = CashDB.add_debt(userid[0][0], amount, name, date)
                if result:
                    return redirect('/debts')
                else:
                    error_message = "Ошибка при записи. Попробуйте ввести другие данные"
    return render_template("debts.html", debts=debts_user, current_datetime=current_datetime)


@app.route('/update_debt_status', methods=['POST'])
def update_debt_status():
    deleted_debt_ids = request.form.getlist('debt_ids[]')
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    for debt_id in deleted_debt_ids:
        result = CashDB.delete_debt(debt_id, userid[0][0])
        if not result:
            flash('Ошибка при удалении долга')
    return redirect('/debts')


@app.route('/wish-list', methods=['POST', 'GET'])
def wish_list():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_wishes = CashDB.show_wishes(userid[0][0])
    current_datetime = datetime.now()
    return render_template("wish-list.html", wishes=user_wishes, current_datetime=current_datetime)


@app.route('/wish-list/add-new-wish', methods=['POST', 'GET'])
def add_new_wish():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_wishes = CashDB.show_wishes(userid[0][0])
    if request.method == 'POST':
        name_wish = request.form['name-wish']
        cost = request.form['cost']
        image = request.files['image']
        date_wish = request.form['date-wish']
        if image.filename != '':
            filename = secure_filename(image.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            image.save(file_path)
            img = Image.open(file_path)
            new_width = 100
            new_height = 100
            img = img.resize((new_width, new_height))
            img.save(file_path)
        result = CashDB.add_wish(cost, userid[0][0], date_wish, name_wish, filename)
        return redirect('/wish-list')
    return render_template("add-new-wish.html")


@app.route('/wish-list/add-money-to-wish', methods=['POST', 'GET'])
def add_money_to_wish():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_wishes = CashDB.show_wishes(userid[0][0])
    if request.method == 'POST':
        selected_wish = request.form['selected_wish']
        print(selected_wish)
        amount = float(request.form['amount'])
        wish_id = CashDB.show_wish_id(selected_wish, userid[0][0])
        primary_amount = CashDB.show_wishes_amount(wish_id[0][0])
        final_amount = float(primary_amount[0][0]) + amount
        result_to_add_amount = CashDB.add_amount_to_wish(final_amount, wish_id[0][0])
        print(result_to_add_amount)
        return redirect('/wish-list')
    return render_template("add-money-to-wish.html", wishes=user_wishes)


@app.route('/static/images_wishes/<filename>')
def uploaded_image(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


@app.route('/bills')
def bills():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_bills = CashDB.show_bills(userid[0][0])
    return render_template("bills.html", bills=user_bills)


@app.route('/bills/add-new-bill', methods=['POST', 'GET'])
def add_new_bill():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    if request.method == 'POST':
        name_bill = request.form['name-bill']
        amount = request.form['amount']
        result = CashDB.add_bill(name_bill, amount, userid[0][0])
        return redirect('/bills')
    return render_template("add-new-bill.html")


@app.route('/bills/add-money-to-bill', methods=['POST', 'GET'])
def add_money_to_bill():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_bills = CashDB.show_bills(userid[0][0])
    if request.method == 'POST':
        selected_bill = request.form['selected_bill']
        balance = float(request.form['amount'])
        bill_id = CashDB.show_bill_id(selected_bill, userid[0][0])
        primary_amount = CashDB.show_bill_balance(bill_id[0][0])
        final_amount = float(primary_amount[0][0]) + balance
        result_to_add_amount = CashDB.add_amount_to_bill(final_amount, bill_id[0][0])
        return redirect('/bills')
    return render_template("add-money-to-bill.html", bills=user_bills)


@app.route('/bills/remove-bill', methods=['POST', 'GET'])
def remove_bill():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_bills = CashDB.show_bills(userid[0][0])
    error_message_transfer = None
    error_message_removing = None
    if request.method == 'POST':
        selected_bill_remove = request.form['selected_bill_remove']
        selected_bill_to_move_money = request.form['selected_bill_to_move_money']
        id_bill_to_remove = CashDB.show_bill_id(selected_bill_remove, userid[0][0])
        id_bill_to_transfer_money = CashDB.show_bill_id(selected_bill_to_move_money, userid[0][0])
        balance = float(CashDB.show_bill_balance(id_bill_to_remove[0][0])[0][0])
        result_transfer = CashDB.transfer_amount_between_bills(id_bill_to_remove[0][0], id_bill_to_transfer_money[0][0], balance)
        if not result_transfer:
            error_message_transfer = "Ошибка при попытке перевода. Попробуйте снова"
            return render_template("remove-bill.html", bills=user_bills, error_message_transfer=error_message_transfer)
        result_removing = CashDB.remove_bill(id_bill_to_remove[0][0], userid[0][0])
        if not result_removing:
            error_message_removing = "Ошибка при попытке удаления счета. Попробуйте снова"
            return render_template("remove-bill.html", bills=user_bills, error_message_removing=error_message_removing)
        return redirect('/bills')
    return render_template("remove-bill.html", bills=user_bills)


@app.route('/bills/transfer', methods=['POST', 'GET'])
def transfer():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_bills = CashDB.show_bills(userid[0][0])
    if request.method == 'POST':
        selected_bill_first = request.form['selected_bill_first']
        selected_bill_second = request.form['selected_bill_second']
        amount = float(request.form['amount'])
        id_bill_first = CashDB.show_bill_id(selected_bill_first, userid[0][0])
        id_bill_second = CashDB.show_bill_id(selected_bill_second, userid[0][0])
        CashDB.transfer_amount_between_bills(id_bill_first[0][0], id_bill_second[0][0], amount)
        return redirect('/bills')
    return render_template("transfer_from_bill_to_bill.html", bills=user_bills)


@app.route('/wish-list/add-money-to-wish-from-bill', methods=['POST', 'GET'])
def add_money_to_wish_from_bill():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_bills = CashDB.show_bills(userid[0][0])
    user_wishes = CashDB.show_wishes(userid[0][0])
    if request.method == 'POST':
        selected_wish = request.form['selected_wish']
        selected_bill = request.form['selected_bill']
        amount_to_add = float(request.form['amount-to-add'])
        id_wish = CashDB.show_wish_id(selected_wish, userid[0][0])
        id_bill = CashDB.show_bill_id(selected_bill, userid[0][0])
        CashDB.transfer_amount_to_wish(id_wish[0][0], id_bill[0][0], amount_to_add)
        return redirect('/wish-list')
    return render_template("add-money-to-wish-from-bill.html", bills=user_bills, wishes=user_wishes)


@app.route('/transaction')
def transaction():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    user_transactions = CashDB.show_transactions(userid[0][0])
    return render_template("transaction.html", transactions=user_transactions, cash_db=CashDB)


@app.route('/transaction/search', methods=['POST', 'GET'])
def search():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    if request.method == 'POST':
        date1 = request.form['date1']
        date2 = request.form['date2']
        result = CashDB.search(date1, date2, userid[0][0])
        return render_template("transaction.html", transactions=result, cash_db=CashDB)
    return render_template("transaction.html")


@app.route('/transaction/add-income', methods=['POST', 'GET'])
def add_income():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    if request.method == 'POST':
        amount = request.form['amount']
        result = CashDB.add_income(amount, userid[0][0], "Доход")
        return redirect('/transaction')
    return render_template("add-income.html")


@app.route('/transaction/add-expense', methods=['POST', 'GET'])
def add_expense():
    email_user = session['email']
    userid = CashDB.get_user_id(email_user)
    categories = CashDB.show_categories()
    if request.method == 'POST':
        amount = request.form['amount']
        choosen_category = request.form['selected_category']
        id_category = CashDB.get_category_id(choosen_category)
        result = CashDB.add_expense(amount, userid[0][0], 'Расход', id_category[0][0])
        return redirect('/transaction')
    return render_template("add-expense.html", categories=categories)

@app.route('/logout')
def logout():
    session.pop('email', None)
    return redirect('/')


@app.route('/debts/download_report', methods=['POST', 'GET'])
def download_report():
    if request.method == "POST":
        email_user = session['email']
        userid = CashDB.get_user_id(email_user)
        doc = Document()
        doc.add_heading('Отчет о долгах', level=1)
        doc.styles['Heading 1'].font.size = Pt(14)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = '#'
        header_cells[1].text = 'Сумма'
        header_cells[2].text = 'Имя или название организации'
        header_cells[3].text = 'Дата возврата'
        debts = CashDB.show_debts(userid[0][0])
        index = 1
        for debt in debts:
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = str(debt[2])
            row_cells[2].text = str(debt[3])
            if debt[4] is not None:
                row_cells[3].text = debt[4].strftime('%d %B %Y')
            index +=1
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers["Content-Disposition"] = "attachment; filename=report.docx"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return response


@app.route('/statistic', methods=['POST', 'GET'])
def statistic():
    if request.method == 'POST':
        email_user = session['email']
        userid = CashDB.get_user_id(email_user)
        chosen_interval = request.form['selected_interval']
        type_stat = request.form['type_stat']
        if chosen_interval == 'За последнюю неделю':
            x_labels = pd.date_range(end=pd.Timestamp.today(), periods=7, freq='D').strftime('%d %b')
        elif chosen_interval == 'За последних 3 месяца':
            x_labels = pd.date_range(end=pd.Timestamp.today(), periods=3, freq='M').strftime('%B')
        elif chosen_interval == 'За последние 3 года':
            x_labels = pd.date_range(end=pd.Timestamp.today(), periods=3, freq='Y').strftime('%Y')
        else:
            return "Выбран недопустимый интервал"


        grouped_stats = None
        if chosen_interval == 'За последнюю неделю':
            today = pd.Timestamp.today().normalize()
            last_week = today - pd.DateOffset(weeks=1)
            last_week_data = CashDB.get_data_statistic_interval(userid[0][0], type_stat, last_week, today)
            df = pd.DataFrame(last_week_data, columns=['amount', 'date'])
            df['date'] = pd.to_datetime(df['date'])
            df = df.sort_values(by='date', ascending=True)
            grouped_stats = df.groupby(df['date'].dt.date)['amount'].sum()
            x_labels = pd.to_datetime(grouped_stats.index).strftime('%d %b')
        elif chosen_interval == 'За последних 3 месяца':
            today = pd.Timestamp.today().normalize()
            last_month = today - pd.DateOffset(months=3)
            last_month_data = CashDB.get_data_statistic_interval(userid[0][0], type_stat, last_month, today)
            df = pd.DataFrame(last_month_data, columns=['amount', 'date'])
            df['date'] = pd.to_datetime(df['date'])
            df = df.sort_values(by='date', ascending=True)
            grouped_stats = df.groupby(df['date'].dt.month)['amount'].sum()
            x_labels = pd.to_datetime(grouped_stats.index, format='%m').strftime('%b')
        elif chosen_interval == 'За последние 3 года':
            today = pd.Timestamp.today().normalize()
            last_year = today - pd.DateOffset(years=3)
            last_year_data = CashDB.get_data_statistic_interval(userid[0][0],type_stat, last_year, today)
            df = pd.DataFrame(last_year_data, columns=['amount', 'date'])
            df['date'] = pd.to_datetime(df['date'])
            df = df.sort_values(by='date', ascending=True)
            grouped_stats = df.groupby(df['date'].dt.year)['amount'].sum()
            x_labels = pd.to_datetime(grouped_stats.index, format='%Y').strftime('%Y')

        if grouped_stats is not None:
            x = np.arange(len(grouped_stats))
            y = grouped_stats.values
            plt.figure()
            plt.plot(x, y, color='red')
            plt.xlabel('День' if chosen_interval == 'За последнюю неделю' else 'Месяц' if chosen_interval == 'За последних 3 месяца' else 'Год')
            plt.ylabel('Сумма')
            plt.title(
                'Статистика за неделю: {}'.format(type_stat) if chosen_interval == 'За последнюю неделю' else 'Статистика по месяцам: {}'.format(type_stat) if chosen_interval == 'За последних 3 месяца' else 'Статистика по годам: {}'.format(type_stat))
            plt.xticks(range(len(x_labels)), x_labels)
            image_stream = io.BytesIO()
            plt.savefig(image_stream, format='png')
            image_stream.seek(0)
            image_base64 = base64.b64encode(image_stream.read()).decode('utf-8')
        return render_template("statistic.html", image_base64=image_base64)
    return render_template("statistic.html")



@app.route('/admin-page', methods=['POST', 'GET'])
def admin_page():
    if 'email' in session:
        email = session['email']
        return render_template("admin-page.html")
    else:
        return redirect(url_for('sign_in'))

@app.route('/admin-info-users', methods=['POST', 'GET'])
def admin_info_users():
    users = CashDB.get_all_users()
    return render_template("admin-info-users.html", users=users)


@app.route('/admin-add-categories', methods=['POST', 'GET'])
def add_new_category():
    all_categories = CashDB.get_all_categories()
    if request.method == 'POST':
        name_category = request.form['name_category']
        for category in all_categories:
            if name_category == category[0]:
                error_message = "Ошибка! Данная категория уже была добавлена."
                return render_template("admin-add-categories.html", categories=all_categories, error_message=error_message)
        result = CashDB.add_categories(name_category)
        return redirect('/admin-page')
    return render_template("admin-add-categories.html", categories=all_categories)


@app.route('/admin-info-users/user-info', methods=['POST', 'GET'])
def user_info():
    email = request.form['selected_user']
    print(email)
    users = CashDB.get_all_users()
    if request.method == 'POST':
        for user in users:
            if user[3] == email:
                name = user[1]
                surname = user[2]
                user_id = user[0]
                debts = CashDB.show_debts(user_id)
                wishes = CashDB.show_wishes(user_id)
                transactions = CashDB.show_transactions(user_id)
                bills = CashDB.show_bills(user_id)
                return render_template("user-info.html", email=email, name_user=name, surname_user=surname, debts=debts, wishes=wishes, transactions=transactions, bills=bills, cash_db=CashDB)
    return render_template("user-info.html")

@app.route('/export-import-json', methods=['GET', 'POST'])
def export_import_page():
    message = request.args.get('message')
    error_message = request.args.get('error_message')
    return render_template('export-import-json.html', message=message, error_message=error_message)

@app.route('/export_data', methods=['GET', 'POST'])
def export_data():
    if request.method == "GET" or request.method == "POST":
        debts = CashDB.debts_table()
        users = CashDB.users_table()
        wishlists = CashDB.wishlists_table()
        bills = CashDB.bills_table()
        system_admin = CashDB.system_admin_table()
        categoriesexpence = CashDB.categoriesexpence_table()
        transaction_money = CashDB.transaction_money_table()

        debts_list = []
        for debt_item in debts:
            debts_list.append({
                "debt_id": str(debt_item[0]),
                "user_id": str(debt_item[1]),
                "debt_amount": float(debt_item[2]),
                "creditor_name": debt_item[3],
                "date": debt_item[4].strftime("%Y-%m-%d")
            })
        with open("debts.json", "w") as json_file:
            json.dump(debts_list, json_file, default=float)

        users_list = []
        for user in users:
            users_list.append({
                "user_id": user[0],
                "firstname": str(user[1]),
                "lastname": str(user[2]),
                "email": user[3],
                "password": str(user[4].tobytes())
            })
        with open("users.json", "w") as json_file:
            json.dump(users_list, json_file)

        wishlists_list = []
        for wish in wishlists:
            wishlists_list.append({
                "wishid": str(wish[0]),
                "purchasedate": wish[1].strftime("%Y-%m-%d"),
                "cost": float(wish[2]),
                "userid": wish[3],
                "itemname": str(wish[4]),
                "targetdate": wish[5].strftime("%Y-%m-%d"),
                "imageurl": wish[6],
                "accummoney": float(wish[7]),
            })
        with open("wishlists.json", "w") as json_file:
            json.dump(wishlists_list, json_file, default=float)

        bills_list = []
        for bill in bills:
            bills_list.append({
                "billid": bill[0],
                "billname": bill[1],
                "balance": float(bill[2]),
                "opendate": bill[3].strftime("%Y-%m-%d"),
                "userid": bill[4]
            })
        with open("bills.json", "w") as json_file:
            json.dump(bills_list, json_file, default=float)

        system_admin_list = []
        for admin in system_admin:
            system_admin_list.append({
                "adminid": admin[0],
                "firstname": admin[1],
                "lastname": admin[2],
                "adminpassword": str(admin[3].tobytes()),
                "adminemail": admin[4]
            })
        with open("system_admin.json", "w") as json_file:
            json.dump(system_admin_list, json_file)

        categoriesexpence_list = []

        for categ in categoriesexpence:
            categoriesexpence_list.append({
                "categoryid": categ[0],
                "categoryname": str(categ[1])
            })
        with open("categoriesexpence.json", "w") as json_file:
            json.dump(categoriesexpence_list, json_file)


        transaction_money_list = []

        for transaction in transaction_money:
            transaction_money_list.append({
                "transactionid": transaction[0],
                "userid": transaction[1],
                "amount": float(transaction[2]),
                "trdate": transaction[3].strftime("%Y-%m-%d"),
                "trtype": str(transaction[4]),
                "categoryid": transaction[5]
            })
        with open("transaction_money.json", "w") as json_file:
            json.dump(transaction_money_list, json_file)

    return redirect('/export-import-json?message=Экспорт%20данных%20прошел%20успешно')



user_schema = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "user_id": {"type": "integer"},
            "firstname": {"type": "string"},
            "lastname": {"type": "string"},
            "email": {"type": "string"},
            "password": {"type": "string"}
        },
        "required": ["user_id", "firstname", "lastname", "email", "password"]
    }
}


import json

@app.route('/import', methods=['POST'])
def import_data():
    if 'import' in request.files:
        json_file = request.files['import']
        if json_file.filename.endswith('.json'):
            try:
                json_data = json_file.read().decode('utf-8')
                data = json.loads(json_data)
                if isinstance(data, dict):
                    # Если в JSON файле только одна запись, преобразуем её в список
                    data = [data]
                jsonschema.validate(data, user_schema)

                for user_data in data:
                    # Обработка данных пользователя
                    user_id = int(user_data['user_id'])
                    user_first_name = user_data['firstname']
                    user_last_name = user_data['lastname']
                    email = user_data['email']
                    password = user_data['password']
                    hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
                    # Добавление пользователя в базу данных
                    if CashDB.add_user_with_id(user_id, user_first_name, user_last_name, email, hashed_password):
                        app.logger.info(f"Пользователь с ID {user_id} успешно добавлен.")
                    else:
                        error_message = f"Ошибка при добавлении пользователя с ID {user_id}."
                        app.logger.error(error_message)

                message = "Данные пользователей успешно импортированы в базу данных."
                return redirect(url_for('export_import_page', message=message))
            except jsonschema.ValidationError as e:
                error_message = "Ошибка валидации JSON"
                print(e)
                return redirect(url_for('export_import_page', error_message=error_message))
            except json.JSONDecodeError as e:
                error_message = "Ошибка декодирования JSON"
                print(e)
                return redirect(url_for('export_import_page', error_message=error_message))
        else:
            error_message = "Неверное расширение файла. Ожидается файл с расширением .json."
            return redirect(url_for('export_import_page', error_message=error_message))
    else:
        error_message = "Ошибка при импорте данных. Файл не найден."
        return redirect(url_for('export_import_page', error_message=error_message))


if __name__ == "__main__":
    app.run(debug=True)
